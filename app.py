import nltk
nltk.download('stopwords')
import streamlit as st
import os
from pyresparser import ResumeParser
from pyresparser import utils
import re # Using this later
import pandas as pd
import PyPDF2
import pdfplumber
from docx import Document
from stqdm import stqdm
from random import randint

st.set_page_config(page_title="Resume Extractor", page_icon=":tired:", layout="wide")

# -- Header Section --
with st.container():
    st.title("Resume Parser :page_facing_up:")

with st.container():

    def single_extract(file):
        # df = pd.DataFrame({'Name':[],'Phone':[],'Email':[],'Skills':[]})
        phone_list = []
        try:
            if file.name.endswith('pdf'):
            # pdfFileObj = open(file,'rb')
            # pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
                with pdfplumber.open(file) as pdf:
                    # for i in range(pdfReader.numPages):
                    page = pdf.pages[0]
                    text = page.extract_text()
                    for i in re.findall(r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]', text):
                        if i in re.findall(r'2[0-9][0-9][0-9][\s\S\D][\s\S\D][\s\S\D]2[0-9][0-9][0-9]', text):
                            continue
                        else:
                            phone_list.append(i)
            elif file.name.endswith('docx'):
                doc = Document(file)
                for para in range(5):
                    for i in re.findall(r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]', doc.paragraphs[para].text):
                        if i in re.findall(r'2[0-9][0-9][0-9][\s\S\D][\s\S\D][\s\S\D]2[0-9][0-9][0-9]', doc.paragraphs[para].text):
                            continue
                        else:
                            phone_list.append(i)
            if phone_list == []:
                phone_list=None
            filename_temp = file.name
            filename_temp = filename_temp.replace('.', ' ')
            filename_temp = filename_temp.replace('_', ' ')
            filename_temp = filename_temp.replace('-', ' ')

            filename_temp = filename_temp.split()[0:5]
            unwanted_list = ['pdf','docx','cv','curriculum','vitae','resume','-','jan','feb','mar','apr','may','jun','jul','aug','sep','oct','nov','dec','2020','2021','2022','2023','2024','2025']
            name = [s for s in filename_temp if s.lower() not in unwanted_list]

            name = ' '.join(name)

            data = ResumeParser(file).get_extracted_data()

            email = data['email']
            skills = data['skills']

            name_dict = {'Name':[name],'Phone':[phone_list],'Email':[email],'Skills':[skills]}
            df = pd.DataFrame(name_dict)
            
        except Exception:
            print("Data encountered an error")
            pass

        # df = pd.concat([df,temp_df],ignore_index=True)

        return(df)
    
    st.write('---')
    with st.form("my-form", clear_on_submit=True):
        uploaded_files = st.file_uploader("Choose a file (Only pdf and docx formats accepted)",type=['pdf','docx'], accept_multiple_files=True)
        submitted = st.form_submit_button("Process Files")

        if submitted and uploaded_files is not None:
            df = pd.DataFrame({'Name':[],'Phone':[],'Email':[],'Skills':[]})
            for uploaded_file, percent_complete in zip(uploaded_files,stqdm(range(len(uploaded_files)))):
                df = pd.concat([df,single_extract(uploaded_file)],ignore_index=True)
            st.table(df)
            
            
            
            
with st.container():

    def single_extract(file):
        # df = pd.DataFrame({'Name':[],'Phone':[],'Email':[],'Skills':[]})
        phone_list = []
        try:
            if file.name.endswith('pdf'):
            # pdfFileObj = open(file,'rb')
            # pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
                with pdfplumber.open(file) as pdf:
                    # for i in range(pdfReader.numPages):
                    page = pdf.pages[0]
                    text = page.extract_text()
                    for i in re.findall(r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]', text):
                        if i in re.findall(r'2[0-9][0-9][0-9][\s\S\D][\s\S\D][\s\S\D]2[0-9][0-9][0-9]', text):
                            continue
                        else:
                            phone_list.append(i)
            elif file.name.endswith('docx'):
                doc = Document(file)
                for para in range(5):
                    for i in re.findall(r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]', doc.paragraphs[para].text):
                        if i in re.findall(r'2[0-9][0-9][0-9][\s\S\D][\s\S\D][\s\S\D]2[0-9][0-9][0-9]', doc.paragraphs[para].text):
                            continue
                        else:
                            phone_list.append(i)
            if phone_list == []:
                phone_list=None
            filename_temp = file.name
            filename_temp = filename_temp.replace('.', ' ')
            filename_temp = filename_temp.replace('_', ' ')
            filename_temp = filename_temp.replace('-', ' ')

            filename_temp = filename_temp.split()[0:5]
            unwanted_list = ['pdf','docx','cv','curriculum','vitae','resume','-','jan','feb','mar','apr','may','jun','jul','aug','sep','oct','nov','dec','2020','2021','2022','2023','2024','2025']
            name = [s for s in filename_temp if s.lower() not in unwanted_list]

            name = ' '.join(name)

            data = ResumeParser(file).get_extracted_data()

            email = data['email']
            skills = data['skills']

            name_dict = {'Name':[name],'Phone':[phone_list],'Email':[email],'Skills':[skills]}
            df = pd.DataFrame(name_dict)
            
        except Exception:
            print("Data encountered an error")
            pass


        return(df)
    st.write('---')
    st.header('If you want to mass download csv files use upload provided below')
    uploaded_files = st.file_uploader("Choose a file (Only pdf and docx formats accepted)",type=['pdf','docx'], accept_multiple_files=True)
    if uploaded_files is not None:
        # file_details = {'filename':uploaded_file.name, 'filetype':uploaded_file.type}
        # st.write(file_details)

        df = pd.DataFrame({'Name':[],'Phone':[],'Email':[],'Skills':[]})
        for uploaded_file, percent_complete in zip(uploaded_files,stqdm(range(len(uploaded_files)))):
            df = pd.concat([df,single_extract(uploaded_file)],ignore_index=True)
        st.table(df)

    @st.cache
    def convert_df(df):
    # IMPORTANT: Cache the conversion to prevent computation on every rerun
        return df.to_csv().encode('utf-8')

    csv = convert_df(df)

    st.download_button(
        label="Download data as CSV",
        data=csv,
        file_name='resume_data.csv',
        mime='text/csv',
    )
    
